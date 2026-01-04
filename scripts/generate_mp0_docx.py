from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

# Paths
BASE_DIR = r"d:\gptWebapp\new-lan-learning\language-learning2"
ARTIFACTS_DIR = r"C:\Users\92515\.gemini\antigravity\brain\cb62c7b0-1621-42e3-98f4-65dbf0c95631"
MARKDOWN_FILE = os.path.join(BASE_DIR, "docs", "paper", "MP0_Draft_v1.md")
OUTPUT_FILE = os.path.join(BASE_DIR, "docs", "paper", "MP0_Final.docx")

# Regex pattern for inline formatting
INLINE_PATTERN = r'(`[^`]+`)|(\$[^$]+\$)|(\*\*[^*]+\*\*)|(\*[^*]+\*)|(\[[^\]]+\]\([^)]+\))'

def apply_inline_formatting(paragraph, text):
    """Apply inline formatting to a paragraph based on Markdown syntax."""
    tokens = re.split(INLINE_PATTERN, text)
    
    for token in tokens:
        if not token: 
            continue
        
        if token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
        elif token.startswith('$') and token.endswith('$'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Cambria Math'
            run.font.italic = True
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('[') and ']' in token and '(' in token and token.endswith(')'):
            # Link: [text](url) -> just text
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', token)
            if m:
                link_text = m.group(1)
                run = paragraph.add_run(link_text)
                run.underline = True
            else:
                paragraph.add_run(token)
        else:
            paragraph.add_run(token)

def create_document():
    doc = Document()
    
    # Read Markdown
    with open(MARKDOWN_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    # Define styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Remove HTML comments
    content = re.sub(r'<!--.*?-->', '', content, flags=re.DOTALL)
    
    lines = content.split('\n')
    
    state = "NORMAL"
    code_buffer = []
    
    for line in lines:
        stripped_line = line.strip()
        
        # --- Code Block Handling ---
        if stripped_line.startswith("```"):
            if state == "NORMAL":
                state = "CODE_BLOCK"
                continue
            elif state == "CODE_BLOCK":
                state = "NORMAL"
                if code_buffer:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(6)
                    
                    full_code = "\n".join(code_buffer)
                    run = p.add_run(full_code)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    
                    code_buffer = []
                continue
        
        if state == "CODE_BLOCK":
            code_buffer.append(line) 
            continue

        # --- Math Block Handling ---
        if stripped_line.startswith("$$") and stripped_line.endswith("$$") and len(stripped_line) > 2:
            formula = stripped_line.strip("$").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.font.name = 'Cambria Math'
            run.font.italic = True
            run.font.size = Pt(11)
            continue

        # --- Normal Text Processing ---
        if not stripped_line:
            continue
            
        if stripped_line.startswith('# '):
            p = doc.add_heading(stripped_line[2:], 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped_line.startswith('## '):
            doc.add_heading(stripped_line[3:], 1)
        elif stripped_line.startswith('### '):
            doc.add_heading(stripped_line[4:], 2)
        elif stripped_line.startswith('#### '):
            doc.add_heading(stripped_line[5:], 3)
        
        elif stripped_line.startswith('**[Figure') or (stripped_line.startswith('*(') and stripped_line.endswith(')*')):
            caption = stripped_line.replace('**', '').strip()
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].italic = True
            
        elif stripped_line.startswith('- ') or stripped_line.startswith('* '):
            # Bullet list with inline formatting
            p = doc.add_paragraph(style='List Bullet')
            apply_inline_formatting(p, stripped_line[2:])
            
        elif len(stripped_line) > 3 and stripped_line[0].isdigit() and stripped_line[1:3] == '. ':
            # Numbered list with inline formatting (e.g., "1. **Bold** text")
            p = doc.add_paragraph(style='List Number')
            # Keep the number prefix
            apply_inline_formatting(p, stripped_line)
            
        else:
            # Standard paragraph with inline formatting
            p = doc.add_paragraph()
            apply_inline_formatting(p, line)

    try:
        doc.save(OUTPUT_FILE)
        print(f"Document saved to {OUTPUT_FILE}")
    except PermissionError:
        print(f"Error: Could not save to {OUTPUT_FILE}. The file might be open.")
        alt_file = OUTPUT_FILE.replace(".docx", "_v5.docx")
        doc.save(alt_file)
        print(f"Document saved to {alt_file} instead.")

if __name__ == "__main__":
    create_document()
