import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    # Force Times New Roman for complex scripts if needed, though mainly for English here
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

def create_document():
    doc = docx.Document()
    
    # Page Setup
    section = doc.sections[0]
    section.top_margin = Inches(1.18)
    section.bottom_margin = Inches(1.18)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)
    
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Read content
    with open('docs/paper/MP0_Final_Content_Strict_EN.txt', 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_buffer = []
    
    for line in lines:
        line = line.strip()
        
        # Skip empty lines unless in code block
        if not line and not in_code_block:
            continue
            
        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run('\n'.join(code_buffer))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_buffer = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
            
        if in_code_block:
            code_buffer.append(line)
            continue
            
        # Handle Headers
        if line.startswith('Chapter '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            set_font(run, size=18, bold=True)
            # Add spacing after chapter title
            p.paragraph_format.space_after = Pt(12)
            
        elif re.match(r'^\d+\.\d+\s', line): # 1.1, 2.1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            set_font(run, size=12, bold=True)
            p.paragraph_format.space_before = Pt(12)
            
        elif re.match(r'^\d+\.\d+\.\d+\s', line): # 1.1.1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            set_font(run, size=12, bold=False) # Normal style as requested
            
        elif line == 'Bibliography' or line == 'Appendix' or line == 'Table of Contents':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(line)
            set_font(run, size=18, bold=True)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            
        elif line.startswith('Design and Development'): # Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_font(run, size=18, bold=True)
            p.paragraph_format.space_after = Pt(24)
            
        elif line.startswith('Master Report') or line.startswith('In partial fulfillment'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_font(run, size=12)
            
        elif line.startswith('December') or line.startswith('The Kyoto College') or line.startswith('Applied Information') or line.startswith('Web Business'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            set_font(run, size=12)
            
        elif line.startswith('Yang Jizhou') or line.startswith('M24W0470') or line.startswith('YANG JIZHOU'):
             # Cover page name section
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            if 'Yang Jizhou' in line: # Name
                 set_font(run, size=18)
            elif 'M24W0470' in line: # ID
                 set_font(run, size=14)
            else:
                 set_font(run, size=12)

        else:
            # Normal text
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run(line)
            set_font(run, size=12)
            
    # Add Page Numbers (Footer)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p.add_run())
    
    doc.save('docs/paper/MP0_Final_Generated.docx')
    print("Document generated successfully.")

if __name__ == '__main__':
    create_document()
