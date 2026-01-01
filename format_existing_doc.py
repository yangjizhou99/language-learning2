import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)

def format_document():
    doc_path = 'docs/paper/MP0.docx'
    try:
        doc = docx.Document(doc_path)
    except Exception as e:
        print(f"Error opening file: {e}")
        return

    # 1. Page Setup (Margins)
    for section in doc.sections:
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(1.18)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)
        
        # Footer (Page Numbers)
        footer = section.footer
        # Clear existing footer paragraphs
        for p in footer.paragraphs:
            p.clear()
        # Add new page number
        if not footer.paragraphs:
            footer.add_paragraph()
        add_page_number(footer.paragraphs[0])

    in_appendix = False

    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Reset paragraph formatting
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        # Check for Appendix start
        if text == 'Appendix':
            in_appendix = True
            
        # --- HEADER LOGIC ---
        if text.startswith('Chapter ') or text in ['Bibliography', 'Appendix', 'Table of Contents']:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            for run in para.runs:
                set_font(run, size=18, bold=True)
                
        elif re.match(r'^\d+\.\d+\s', text) and not re.match(r'^\d+\.\d+\.\d+', text): # 1.1, 2.1 (Level 2)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            for run in para.runs:
                set_font(run, size=12, bold=True)
                
        elif re.match(r'^\d+\.\d+\.\d+\s', text): # 1.1.1 (Level 3 - Normal)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(6)
            for run in para.runs:
                set_font(run, size=12, bold=False)

        # --- TITLE PAGE LOGIC (Heuristic) ---
        elif any(x in text for x in ['Design and Development', 'Master Report', 'In partial fulfillment', 'The Kyoto College', 'Applied Information', 'Web Business', 'Yang Jizhou', 'M24W0470', 'YANG JIZHOU', 'December 28']):
             para.alignment = WD_ALIGN_PARAGRAPH.CENTER
             for run in para.runs:
                 if 'Design and Development' in text or 'Yang Jizhou' in text: # Main Title or Name
                     set_font(run, size=18, bold=('Design' in text))
                 elif 'M24W0470' in text or 'Date' in text:
                     set_font(run, size=14)
                 else:
                     set_font(run, size=12)

        # --- APPENDIX / CODE LOGIC ---
        elif in_appendix and (text.startswith('import') or text.startswith('export') or '{' in text or '}' in text or ';' in text or 'return' in text):
             para.alignment = WD_ALIGN_PARAGRAPH.LEFT
             for run in para.runs:
                 set_font(run, font_name='Courier New', size=10)

        # --- NORMAL TEXT ---
        else:
            if not in_appendix:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT # Code/Appendix text left aligned
                
            for run in para.runs:
                # Preserve bold if it was already bold (e.g. in-text emphasis), but force font
                is_bold = run.bold
                if in_appendix:
                     set_font(run, font_name='Courier New', size=10, bold=is_bold)
                else:
                     set_font(run, font_name='Times New Roman', size=12, bold=is_bold)

    doc.save('docs/paper/MP0_Formatted.docx')
    print("Formatted document saved as docs/paper/MP0_Formatted.docx")

if __name__ == '__main__':
    format_document()
